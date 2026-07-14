from __future__ import annotations

import os
import zipfile
from collections import Counter

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


DOCX_PATH = r"D:\我的专利（1）\ALNS说明书 .docx"


def text_of(paragraph: Paragraph) -> str:
    return paragraph.text.strip().replace("\u3000", " ")


def set_text(paragraph: Paragraph, text: str) -> None:
    element = paragraph._element
    for child in list(element):
        element.remove(child)
    if text:
        paragraph.add_run(text)


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if paragraph.style is not None:
        new_para.style = paragraph.style
    new_para.add_run(text)
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def find_index(paragraphs: list[Paragraph], predicate) -> int:
    for index, paragraph in enumerate(paragraphs):
        if predicate(text_of(paragraph)):
            return index
    raise ValueError("Expected paragraph was not found.")


def replace_invention_steps(doc: Document) -> None:
    paragraphs = doc.paragraphs
    start = find_index(paragraphs, lambda t: t.startswith("[0008] 步骤一："))
    end = find_index(paragraphs, lambda t: t.startswith("[0015] 进一步地"))

    anchor = paragraphs[start - 1]
    new_steps = [
        "[0008] 步骤一：根据配送任务、无人机、无人车以及配送环境信息，建立空地协同配送基础数据模型。",
        "[0009] 步骤二：针对每个任务与无人机组合构建候选服务方式集合，所述候选服务方式至少包括无人机直接配送方式和无人机与无人车接力配送方式。",
        "[0010] 步骤三：对所述候选服务方式进行可行性筛选，筛除不满足位置有效性、剩余续航、配送完成时间和接力部署条件的候选方案。",
        "[0011] 步骤四：构建统一评价模型，对直接配送方式和接力配送方式分别计算时间成本、无人机能耗、无人车能耗、等待惩罚和回退风险，得到对应候选方案的综合成本。",
        "[0012] 步骤五：基于所述候选方案综合成本，采用启发式插入方式生成初始解，并通过ALNS中的破坏算子与修复算子对当前解进行迭代优化，获得任务分配、设备匹配、路线插入位置和中继点选择的联合优化结果。",
        "[0013] 步骤六：在执行过程中，当接力方案出现部署不可行、同步偏差过大、等待超时或剩余裕度不足时，触发基于同一评价器的回退或重选机制，将当前任务切换为新的接力方案或直接配送方案。",
        "[0014] 步骤七：输出满足约束条件的无人机无人车协同配送方案，所述配送方案至少包括任务分配结果、直接配送或接力配送模式、中继点位置、设备执行顺序以及能耗与时序评估结果。",
        "在步骤一中，具体包括以下三个步骤：",
        "步骤11，采集任务起点、任务终点、载荷、优先级、截止时间、无人机当前位置、无人机剩余电量、无人车当前位置以及配送环境约束信息。",
        "步骤12，依据采集到的任务信息、设备信息和环境信息，建立任务集合、无人机集合、无人车集合以及环境约束集合。",
        "步骤13，对缺失数据、越界坐标、不可执行设备和不满足基本服务条件的任务进行预处理，形成后续候选生成阶段可直接调用的标准化输入。",
        "在步骤二中，具体包括以下三个步骤：",
        "步骤21，为每一个任务与无人机组合生成无人机直接配送候选方案。",
        "步骤22，基于无人车当前位置、任务起点、任务终点、任务走廊采样点以及无人车到任务走廊的投影点生成接力配送候选中继点。",
        "步骤23，对重复中继点、落入禁飞区或障碍区域的中继点以及不满足接力部署条件的中继点进行剔除，得到候选服务方式集合。",
        "在步骤三中，具体包括以下四个步骤：",
        "步骤31，判断候选中继点是否位于有效可执行区域内，若否，则剔除该候选方案。",
        "步骤32，判断无人机是否能够从当前位置连续飞行至候选中继点，若否，则剔除该候选方案。",
        "步骤33，判断无人机到达候选中继点后的剩余航程是否足以完成配送段，若否，则剔除该候选方案。",
        "步骤34，判断候选方案的预计完成时间是否满足任务截止时间要求，若满足，则将该候选方案输入统一评价模型。",
        "在步骤四中，具体包括以下五个步骤：",
        "步骤41，依据候选方案的飞行距离与设备速度计算时间成本。",
        "步骤42，依据无人机飞行距离、无人车行驶距离以及对应单位能耗参数计算能耗成本。",
        "步骤43，依据无人机与无人车到达候选中继点的预计时差计算等待惩罚。",
        "步骤44，依据续航裕度、时间裕度和同步偏差构建回退风险项。",
        "步骤45，对时间成本、无人机能耗、无人车能耗、等待惩罚和回退风险进行加权求和，得到候选方案综合成本。",
        "在步骤五中，具体包括以下五个步骤：",
        "步骤51，按照综合增量成本由低到高的顺序构造初始解。",
        "步骤52，通过随机移除、最差成本移除或风险导向移除执行破坏操作。",
        "步骤53，通过贪婪插入、遗憾值插入或接力感知修复执行恢复操作。",
        "步骤54，结合模拟退火接受准则判断是否接受新解，并更新当前解和历史最优解。",
        "步骤55，当达到最大迭代次数、连续多轮未改进或温度降低至预设下限时，输出最优或近优联合优化结果。",
        "在步骤六中，具体包括以下四个步骤：",
        "步骤61，实时监测无人机剩余电量、无人车到达状态、候选中继点有效性以及任务剩余时间。",
        "步骤62，当等待时间超过阈值、中继点失效、部署不可行或时间裕度不足时，触发异常判定。",
        "步骤63，调用统一评价器对当前可用无人车、候选中继点和直接配送方案进行重新评价。",
        "步骤64，若存在新的可行接力方案，则执行中继点重选；若不存在可行接力方案，则切换为直接配送方案。",
        "在步骤七中，输出结果还包括各任务的预计完成时刻、等待时间、模式风险值以及回退触发记录，用于后续执行校核和调度追溯。",
    ]

    for paragraph in paragraphs[start:end]:
        remove_paragraph(paragraph)

    for line in new_steps:
        anchor = insert_after(anchor, line)


def align_implementation_headings(doc: Document) -> None:
    replacements = {
        "具体实施方法": "具体实施方式",
        "本实施例以园区末端配送场景为例，对本发明的建模对象、候选方案生成、统一评价、ALNS优化、执行控制以及异常回退过程进行详细说明。该实施例对应的系统由任务层、候选生成层、统一评分层、ALNS求解层和执行校核层组成，其中统一评分层同时服务于初始分配、迭代修复和执行阶段回退决策，从而保证评分逻辑与执行逻辑一致。": "本实施例以园区末端配送场景为例，对本发明的建模对象、候选方案生成、统一评价、ALNS优化、执行控制以及异常回退过程进行详细说明。具体实施方式分为六个部分介绍：第一部分为场景建模与基础对象定义，第二部分为候选服务方式生成与可行性筛选，第三部分为统一评价模型，第四部分为基于ALNS的联合优化过程，第五部分为执行阶段的连续接力语义与回退机制，第六部分为结合实际项目参数的实施说明。",
        "一、场景建模与基础对象定义": "第一部分为：场景建模与基础对象定义",
        "二、候选服务方式生成与可行性筛选": "第二部分为：候选服务方式生成与可行性筛选",
        "三、统一评价模型": "第三部分为：统一评价模型",
        "四、基于ALNS的联合优化过程": "第四部分为：基于ALNS的联合优化过程",
        "五、执行阶段的连续接力语义与回退机制": "第五部分为：执行阶段的连续接力语义与回退机制",
        "六、结合实际项目参数的实施说明": "第六部分为：结合实际项目参数的实施说明",
    }
    for paragraph in doc.paragraphs:
        current = text_of(paragraph)
        if current in replacements:
            set_text(paragraph, replacements[current])


def remove_duplicate_explanations(doc: Document) -> int:
    texts = [text_of(p) for p in doc.paragraphs]
    counts = Counter(t for t in texts if t)
    duplicate_candidates = {
        text
        for text, count in counts.items()
        if count > 1
        and (
            text.startswith("上述")
            or text.startswith("在调度决策时")
            or text.startswith("当候选")
            or text.startswith("通过将")
            or text.startswith("其中，第一重")
            or text.startswith("只有满足全部判断条件")
            or text.startswith("在所述公式中")
            or text.startswith("若无人机先")
            or text.startswith("通过设置风险项")
            or text.startswith("因此，所述统一评价器")
            or text.startswith("当某任务的次优插入成本")
            or text.startswith("在温度较高阶段")
        )
    }

    seen: set[str] = set()
    removed = 0
    for paragraph in list(doc.paragraphs):
        current = text_of(paragraph)
        if current in duplicate_candidates:
            if current in seen:
                remove_paragraph(paragraph)
                removed += 1
            else:
                seen.add(current)
    return removed


def validate(path: str) -> dict[str, int | bool]:
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    return {
        "omath_count": xml.count("<m:oMath") + xml.count("<m:oMathPara"),
        "has_raw_latex": "\\" in xml,
        "paragraph_count": xml.count("<w:p"),
    }


def main() -> None:
    if not os.path.exists(DOCX_PATH):
        raise FileNotFoundError(DOCX_PATH)

    before = validate(DOCX_PATH)
    doc = Document(DOCX_PATH)
    replace_invention_steps(doc)
    align_implementation_headings(doc)
    removed = remove_duplicate_explanations(doc)
    doc.save(DOCX_PATH)
    after = validate(DOCX_PATH)
    print({"before": before, "after": after, "removed_duplicates": removed})


if __name__ == "__main__":
    main()
