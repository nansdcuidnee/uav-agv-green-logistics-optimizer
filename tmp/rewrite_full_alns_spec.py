from __future__ import annotations

import os
import re
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


OUTPUT_PATH = r"D:\我的专利（1）\ALNS说明书 .docx"


def set_font(run, name: str = "宋体", size: int = 12, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), name)


def add_paragraph(doc: Document, text: str = "", size: int = 12, bold: bool = False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(3)
    if text:
        run = paragraph.add_run(text)
        set_font(run, size=size, bold=bold)
    return paragraph


def add_heading(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_font(run, size=14, bold=True)
    return paragraph


def math_run(text: str) -> OxmlElement:
    mr = OxmlElement("m:r")
    mt = OxmlElement("m:t")
    mt.text = text
    mr.append(mt)
    return mr


def math_delimiter(elements: list[OxmlElement]) -> OxmlElement:
    node = OxmlElement("m:d")
    node.append(OxmlElement("m:dPr"))
    e = OxmlElement("m:e")
    for element in elements:
        e.append(element)
    node.append(e)
    return node


def math_sub(base: str, sub: str) -> OxmlElement:
    node = OxmlElement("m:sSub")
    e = OxmlElement("m:e")
    e.append(math_run(base))
    s = OxmlElement("m:sub")
    s.append(math_run(sub))
    node.append(e)
    node.append(s)
    return node


def math_sup(base: str, sup: str) -> OxmlElement:
    node = OxmlElement("m:sSup")
    e = OxmlElement("m:e")
    e.append(math_run(base))
    s = OxmlElement("m:sup")
    s.append(math_run(sup))
    node.append(e)
    node.append(s)
    return node


def math_subsup(base: str, sub: str, sup: str) -> OxmlElement:
    node = OxmlElement("m:sSubSup")
    e = OxmlElement("m:e")
    e.append(math_run(base))
    sub_node = OxmlElement("m:sub")
    sub_node.append(math_run(sub))
    sup_node = OxmlElement("m:sup")
    sup_node.append(math_run(sup))
    node.append(e)
    node.append(sub_node)
    node.append(sup_node)
    return node


def math_frac(num_elements: list[OxmlElement], den_elements: list[OxmlElement]) -> OxmlElement:
    node = OxmlElement("m:f")
    node.append(OxmlElement("m:fPr"))
    num = OxmlElement("m:num")
    den = OxmlElement("m:den")
    for element in num_elements:
        num.append(element)
    for element in den_elements:
        den.append(element)
    node.append(num)
    node.append(den)
    return node


def add_equation_elements(doc: Document, elements: list[OxmlElement]):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    omath_para = OxmlElement("m:oMathPara")
    omath = OxmlElement("m:oMath")
    for element in elements:
        omath.append(element)
    omath_para.append(omath)
    paragraph._p.append(omath_para)
    return paragraph


def add_equation(doc: Document, text: str):
    return add_equation_elements(doc, [math_run(text)])


def add_distance_equation(doc: Document, left_sub: str, terms: list[tuple[str, str]]):
    elements: list[OxmlElement] = [math_sub("L", left_sub), math_run("=")]
    for idx, (a, b) in enumerate(terms):
        if idx:
            elements.append(math_run("+"))
        elements.extend([math_run("D"), math_delimiter([math_run(a), math_run(","), math_run(b)])])
    return add_equation_elements(doc, elements)


def add_single_distance_equation(doc: Document, left_sub: str, a: str, b: str):
    return add_equation_elements(
        doc,
        [math_sub("L", left_sub), math_run("="), math_run("D"), math_delimiter([math_run(a), math_run(","), math_run(b)])],
    )


def add_energy_equation(doc: Document, left_sub: str, k_sub: str, l_sub: str):
    return add_equation_elements(
        doc,
        [math_sub("E", left_sub), math_run("="), math_sub("K", k_sub), math_sub("L", l_sub)]
    )


def add_wait_equation(doc: Document):
    return add_equation_elements(
        doc,
        [
            math_run("W=max"),
            math_delimiter([math_sub("T", "uav"), math_run(","), math_sub("T", "agv")]),
            math_run("-min"),
            math_delimiter([math_sub("T", "uav"), math_run(","), math_sub("T", "agv")]),
        ],
    )


def add_risk_equation(doc: Document):
    return add_equation_elements(
        doc,
        [math_run("R=0.4A+0.3B+0.3C")],
    )


def add_cost_equation(doc: Document):
    return add_equation_elements(
        doc,
        [
            math_run("F="),
            math_sub("w", "t"), math_run("T+"),
            math_sub("w", "u"), math_sub("E", "uav"), math_run("+"),
            math_sub("w", "a"), math_sub("E", "agv"), math_run("+"),
            math_sub("w", "c"), math_sub("C", "carbon"), math_run("+"),
            math_sub("w", "w"), math_run("W+"),
            math_sub("w", "r"), math_run("R"),
        ],
    )


def add_corridor_point_equation(doc: Document):
    return add_equation_elements(
        doc,
        [math_sub("P", "k"), math_run("="), math_run("S+"), math_sub("λ", "k"), math_delimiter([math_run("E-S")])],
    )


def add_lambda_equation(doc: Document):
    return add_equation_elements(
        doc,
        [math_sub("λ", "k"), math_run("="), math_frac([math_run("k")], [math_run("m+1")]), math_run("，k=1,2,...,m")],
    )


def add_projection_point_equation(doc: Document):
    return add_equation_elements(
        doc,
        [math_sub("P", "proj"), math_run("="), math_run("S+μ"), math_delimiter([math_run("E-S")])],
    )


def add_projection_mu_equation(doc: Document):
    return add_equation_elements(
        doc,
        [
            math_run("μ=min"),
            math_delimiter([
                math_run("1,max"),
                math_delimiter([
                    math_run("0,"),
                    math_frac(
                        [math_delimiter([math_run("A-S")]), math_run("·"), math_delimiter([math_run("E-S")])],
                        [math_run("D"), math_delimiter([math_run("S"), math_run(","), math_run("E")]), math_run("×"), math_run("D"), math_delimiter([math_run("S"), math_run(","), math_run("E")])],
                    ),
                ]),
            ]),
        ],
    )


def add_available_range_equation(doc: Document):
    return add_equation_elements(
        doc,
        [math_sub("L", "rem"), math_run("="), math_sub("L", "max"), math_run("b-"), math_sub("L", "lock")],
    )


def add_range_constraint_equation(doc: Document):
    return add_equation_elements(
        doc,
        [math_sub("L", "deploy"), math_run("+"), math_sub("L", "service"), math_run("≤"), math_sub("L", "rem")],
    )


def add_arrival_time_equation(doc: Document, left_sub: str, a: str, b: str, v_sub: str):
    return add_equation_elements(
        doc,
        [
            math_sub("T", left_sub), math_run("="), math_run("t0+"),
            math_frac([math_run("D"), math_delimiter([math_run(a), math_run(","), math_run(b)])], [math_sub("v", v_sub)]),
        ],
    )


def add_finish_time_equation(doc: Document):
    return add_equation_elements(
        doc,
        [
            math_sub("T", "finish"), math_run("="), math_sub("T", "start"), math_run("+"),
            math_frac([math_run("D"), math_delimiter([math_run("S"), math_run(","), math_run("E")])], [math_sub("v", "uav")]),
        ],
    )


def add_deadline_constraint_equation(doc: Document):
    return add_equation_elements(doc, [math_sub("T", "finish"), math_run("≤"), math_sub("T", "deadline")])


def add_carbon_equation(doc: Document):
    return add_equation_elements(
        doc,
        [
            math_sub("C", "carbon"), math_run("="),
            math_sub("η", "u"), math_sub("E", "uav"), math_run("+"),
            math_sub("η", "a"), math_sub("E", "agv"),
        ],
    )


def add_insert_cost_equation(doc: Document):
    return add_equation_elements(
        doc,
        [math_sub("C", "insert"), math_run("="), math_sub("C", "after"), math_run("-"), math_sub("C", "before")],
    )


def add_regret_equation(doc: Document):
    return add_equation_elements(
        doc,
        [math_sub("Δ", "j"), math_run("="), math_sub("C", "j,2"), math_run("-"), math_sub("C", "j,1")],
    )


def add_acceptance_equation(doc: Document):
    return add_equation_elements(
        doc,
        [
            math_sub("P", "acc"), math_run("=exp"),
            math_delimiter([
                math_frac([math_sub("F", "old"), math_run("-"), math_sub("F", "new")], [math_run("τ")])
            ]),
        ],
    )


def add_weight_update_equation(doc: Document):
    return add_equation_elements(
        doc,
        [math_sub("Q", "i"), math_run("←"), math_delimiter([math_run("1-ρ")]), math_sub("Q", "i"), math_run("+ρ"), math_sub("S", "i")],
    )


def add_candidate_score_equation(doc: Document):
    return add_equation_elements(
        doc,
        [math_sub("G", "cand"), math_run("="), math_run("F+βW+γR")],
    )


def add_fallback_trigger_equation(doc: Document):
    return add_equation_elements(
        doc,
        [
            math_run("trigger="),
            math_delimiter([
                math_run("b<"), math_sub("b", "min"), math_run(" or W>"), math_sub("W", "max"),
                math_run(" or "), math_sub("M", "time"), math_run("<0"),
            ]),
        ],
    )


def add_reselect_equation(doc: Document):
    return add_equation_elements(
        doc,
        [math_sup("R", "*"), math_run("=argmin"), math_delimiter([math_run("F(R)")])],
    )


def add_title(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(text)
    set_font(run, size=16, bold=True)


def normalize_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_implementation_section(doc: Document) -> None:
    add_heading(doc, "具体实施方式")
    add_paragraph(
        doc,
        "为了对本发明的技术特征、目的和有益效果有更加清楚的理解，现结合本实验项目的实际实现对本发明的技术方案进行详细说明。显然，所描述的实施例是本发明的一部分实施例，而不是全部实施例，不能理解为对本发明可实施范围的限定。基于本发明的实施例，本领域普通技术人员在没有作出创造性劳动前提下所获得的其他实施例，均属于本发明的保护范围。",
    )
    add_paragraph(doc, "实施例一：")
    add_paragraph(doc, "本实施例中，如图1所示，一种基于ALNS的无人机无人车协同绿色配送优化方法，应用于本实验项目中的园区取送货配送场景。系统包括配送任务集合、无人机集合、无人车集合、配送环境对象、中继候选点生成器、统一成本评价器、ALNS破坏修复算子和执行阶段回退模块，包括以下步骤：")
    add_paragraph(doc, "步骤1：获取配送任务、无人机、无人车和环境约束信息，建立空地协同配送基础数据模型。")
    add_paragraph(doc, "步骤2：根据任务与设备状态生成候选服务方式，所述候选服务方式包括无人机直接配送方式和无人机无人车接力配送方式。")
    add_paragraph(doc, "步骤3：对候选服务方式进行可行性筛选，剔除不满足环境有效性、剩余航程、任务截止时间和部署条件的候选方案。")
    add_paragraph(doc, "步骤4：构建统一成本评价模型，对直接配送候选方案和接力配送候选方案计算综合成本、成本分解、预计等待时间、预计时间裕度和模式风险。")
    add_paragraph(doc, "步骤5：根据统一成本评价结果构建候选池，并采用ALNS对任务分配、无人机匹配、无人车匹配、配送模式、路径插入位置和中继点选择进行联合优化。")
    add_paragraph(doc, "步骤6：在任务执行过程中监测接力方案状态，满足回退触发条件时执行中继点重选、接力转直接配送或任务重新挂起。")
    add_paragraph(doc, "步骤7：输出协同配送优化方案，所述方案包括任务分配结果、配送模式、中继点位置、无人机执行序列、无人车移动动作、综合成本和回退事件。")

    add_paragraph(doc, "具体的，所述步骤1具体包括以下子步骤：")
    add_paragraph(doc, "步骤11：构建配送任务集合。系统读取当前调度周期内处于待分配状态的取送货任务，每个任务至少包括任务编号、任务起点、任务终点、载荷、任务状态和截止时间。任务起点表示取货位置，任务终点表示投送位置，截止时间用于计算时间裕度。")
    add_paragraph(doc, "步骤12：构建无人机集合。系统读取处于空闲状态的无人机，每架无人机至少包括当前位置、剩余电量、最大航程、最大飞行速度和当前已规划路线。与将仓库作为统一起点的方式不同，本实施例在直接配送和接力配送评价中均优先使用无人机当前位置作为真实起点。")
    add_paragraph(doc, "步骤13：构建无人车集合。系统读取处于空闲状态的无人车，每辆无人车至少包括当前位置、行驶状态、最大行驶速度和单位距离能耗参数。无人车主要用于到达中继点，为无人机提供接力衔接位置。")
    add_paragraph(doc, "步骤14：构建环境约束集合。环境对象提供位置有效性判断接口，用于判断任务点、中继点和设备位置是否处于有效配送区域内。若候选位置落入障碍区域、禁飞区域或环境边界外，则不进入后续优化过程。")

    add_paragraph(doc, "具体的，所述步骤11包括以下步骤：")
    add_paragraph(doc, "步骤111：读取任务起点和任务终点，形成任务走廊。任务走廊由任务起点到任务终点的线段表示，用于后续生成走廊采样中继点和无人车投影中继点。")
    add_paragraph(doc, "步骤112：读取任务状态，仅将处于待分配状态的任务加入当前调度周期。若任务已经处于执行中或等待无人车状态，则不作为普通新任务重复分配，而是由执行阶段回退模块进行状态检查。")
    add_paragraph(doc, "步骤113：读取任务截止时间。当任务包含截止时间时，统一评价器计算预计完成时间与截止时间之间的时间裕度；当任务不包含截止时间时，系统为该任务设置默认裕度，使其仍可参与统一评价。")

    add_paragraph(doc, "具体的，所述步骤2具体包括以下子步骤：")
    add_paragraph(doc, "步骤21：生成直接配送候选方案。对于每一个任务与无人机组合，系统生成无人机直接配送路径。该路径从无人机当前位置出发，经任务起点到任务终点，并根据路线衔接关系返回当前位置锚点或后续任务衔接点。直接配送飞行距离表示如下：")
    add_distance_equation(doc, "direct", [("U", "S"), ("S", "E"), ("E", "H")])
    add_paragraph(doc, "其中，U表示无人机当前位置，S表示任务起点，E表示任务终点，H表示返回锚点或后续衔接位置，D表示两点之间的距离。该公式体现了本实验项目中直接配送评价使用无人机真实当前位置，而不是默认仓库点作为起点。")
    add_paragraph(doc, "步骤22：生成接力配送候选方案。对于每一个任务、每一架无人机和每一辆空闲无人车，系统生成与该无人车绑定的候选中继点。候选中继点来源包括无人车当前位置、任务起点、任务终点、任务走廊采样点以及无人车当前位置到任务走廊的投影点。")
    add_paragraph(doc, "步骤23：对候选中继点去重。系统将候选点坐标按一位小数进行近似取整，使用取整后的坐标作为去重键，避免同一位置由多个来源重复生成。")

    add_paragraph(doc, "具体的，所述步骤22包括以下步骤：")
    add_paragraph(doc, "步骤221：将无人车当前位置作为候选中继点。该候选点用于保留无人车无需额外移动或少量移动即可参与接力的方案。")
    add_paragraph(doc, "步骤222：将任务起点和任务终点作为候选中继点。该候选点用于覆盖取货点附近接力和投送点附近接力的情况。")
    add_paragraph(doc, "步骤223：生成任务走廊采样点。本实验项目在任务起点与任务终点之间选取25%、50%和75%三个比例位置作为走廊采样点。一般地，第k个走廊采样点可表示为：")
    add_corridor_point_equation(doc)
    add_lambda_equation(doc)
    add_paragraph(doc, "其中，P表示走廊采样点，S表示任务起点，E表示任务终点，λ表示采样比例。在本实施例中，λ依次取0.25、0.50和0.75，以覆盖任务路径前段、中段和后段的接力位置。")
    add_paragraph(doc, "步骤224：生成无人车投影点。系统计算无人车当前位置到任务走廊的投影点，并将投影点限制在任务起点和任务终点形成的线段范围内，投影点表示如下：")
    add_projection_point_equation(doc)
    add_projection_mu_equation(doc)
    add_paragraph(doc, "其中，A表示无人车当前位置，μ表示投影比例。通过投影点可以获得无人车靠近任务走廊的接力位置，使候选点生成更符合实际地面车辆位置。")
    add_paragraph(doc, "步骤225：计算接力配送路径距离。接力配送由无人机部署段和配送段共同组成，其中部署段表示无人机从当前位置飞往候选中继点的距离，表示如下：")
    add_single_distance_equation(doc, "deploy", "U", "R")
    add_paragraph(doc, "无人机到达中继点后，从中继点飞往任务起点，再由任务起点飞往任务终点，最后返回中继点或后续衔接点，配送段距离表示如下：")
    add_distance_equation(doc, "service", [("R", "S"), ("S", "E"), ("E", "R")])
    add_paragraph(doc, "因此，接力配送总飞行距离表示如下：")
    add_distance_equation(doc, "relay", [("U", "R"), ("R", "S"), ("S", "E"), ("E", "R")])
    add_paragraph(doc, "上述公式中的R表示候选中继点。与只计算中继点到任务起点、任务终点的方式不同，本实施例显式计入无人机从当前位置部署到中继点的距离，使接力方案评分更接近实际执行过程。")

    add_paragraph(doc, "具体的，所述步骤3具体包括以下子步骤：")
    add_paragraph(doc, "步骤31：进行环境有效性筛选。系统调用环境对象的位置有效性判断接口，判断候选中继点是否处于可执行配送区域内。若候选中继点无效，则该候选方案直接删除。")
    add_paragraph(doc, "步骤32：进行无人机剩余航程筛选。系统根据无人机最大航程、当前电量比例和已经锁定的路线航程，计算当前可用剩余航程：")
    add_available_range_equation(doc)
    add_paragraph(doc, "其中，Lrem表示剩余可用航程，Lmax表示无人机最大航程，b表示剩余电量比例，Llock表示已经锁定路线占用的航程。接力配送候选方案需要满足如下航程约束：")
    add_range_constraint_equation(doc)
    add_paragraph(doc, "若部署段距离与配送段距离之和大于剩余可用航程，则说明无人机无法安全完成接力配送，该候选方案被剔除。")
    add_paragraph(doc, "步骤33：进行时序可行性筛选。系统分别估计无人机和无人车到达候选中继点的时间。无人机到达中继点时间表示如下：")
    add_arrival_time_equation(doc, "uav", "U", "R", "uav")
    add_paragraph(doc, "无人车到达中继点时间表示如下：")
    add_arrival_time_equation(doc, "agv", "A", "R", "agv")
    add_paragraph(doc, "步骤34：进行任务截止时间筛选。系统根据任务开始执行时间和任务起点到任务终点的配送时间估计任务完成时间：")
    add_finish_time_equation(doc)
    add_deadline_constraint_equation(doc)
    add_paragraph(doc, "若预计完成时间超过任务截止时间，则该候选方案不进入候选池。上述筛选使ALNS搜索阶段只处理基本可行的候选方案，减少无效搜索。")

    add_paragraph(doc, "具体的，所述步骤4具体包括以下子步骤：")
    add_paragraph(doc, "步骤41：计算能耗成本。本实验项目中的统一成本评价器采用简化能耗模型，若能耗模型对象提供单位距离能耗参数，则调用该参数；否则无人机单位距离能耗采用默认值5.0，无人车单位距离能耗采用默认值3.0。能耗表达式为：")
    add_energy_equation(doc, "uav", "uav", "uav")
    add_energy_equation(doc, "agv", "agv", "agv")
    add_paragraph(doc, "其中，E表示能耗，K表示单位距离能耗参数，L表示执行距离。下标uav表示无人机，下标agv表示无人车。")
    add_paragraph(doc, "步骤42：计算碳排放估计值。本实验项目中将无人机能耗和无人车能耗之和乘以折算系数，用于形成绿色配送评价项：")
    add_carbon_equation(doc)
    add_paragraph(doc, "其中，Ccarbon表示碳排放估计值，η表示能耗到碳排放的折算系数。该项用于使算法在时效和能耗之间进行统一权衡。")
    add_paragraph(doc, "步骤43：计算等待同步成本。接力配送需要无人机和无人车在中继点附近形成时间协同，预计等待时间由无人机预计到达时间和无人车预计到达时间之差确定：")
    add_wait_equation(doc)
    add_paragraph(doc, "等待时间越大，说明接力同步性越差。统一评价器将等待时间乘以等待惩罚权重，使同步偏差较大的接力方案排序靠后。")
    add_paragraph(doc, "步骤44：计算模式风险。模式风险由航程裕度风险、时间裕度风险和同步偏差风险组成，表达式为：")
    add_risk_equation(doc)
    add_paragraph(doc, "其中，A表示航程裕度风险项，B表示时间裕度风险项，C表示同步偏差风险项。本实验项目中三项权重分别为0.4、0.3和0.3。")
    add_paragraph(doc, "步骤45：计算综合成本。统一评价器将时间成本、无人机能耗、无人车能耗、碳排放估计、等待惩罚、超时惩罚和回退风险合成为候选方案综合成本：")
    add_cost_equation(doc)
    add_paragraph(doc, "其中，F表示候选方案综合成本，T表示时间成本，W表示等待惩罚，R表示模式风险，w表示各成本项权重。通过该综合成本，直接配送和接力配送能够在同一评价口径下进行比较。")
    add_paragraph(doc, "步骤46：输出评价结果。统一评价器输出可行性标志、成本增量、成本分解、预计等待时间、预计时间裕度、模式风险和对应的配送选项。该评价结果同时用于候选池构建、初始解生成、修复算子插入和执行阶段回退判断。")

    add_paragraph(doc, "具体的，所述步骤5具体包括以下子步骤：")
    add_paragraph(doc, "步骤51：构建候选池。系统针对每一个任务与无人机组合建立候选池。候选池包括直接配送可行标志和接力候选方案列表。本实验项目中候选池默认采用diverse_topk策略，并保留前5个接力候选方案。候选排序指标为：")
    add_candidate_score_equation(doc)
    add_paragraph(doc, "其中，G表示候选排序指标，F表示综合成本，W表示预计等待时间，R表示模式风险。系统优先选择综合成本低、等待时间短、模式风险小的候选方案，同时优先保持不同无人车来源的多样性。")
    add_paragraph(doc, "步骤52：生成初始解。系统采用遗憾值插入方式生成初始解。对于每个未分配任务，系统收集其在不同无人机、不同配送模式和不同中继点下的可行候选方案，并按照成本增量排序。插入成本表示如下：")
    add_insert_cost_equation(doc)
    add_paragraph(doc, "步骤53：计算遗憾值。对于任务j，系统计算其最优候选方案和次优候选方案之间的成本差，遗憾值表示如下：")
    add_regret_equation(doc)
    add_paragraph(doc, "遗憾值越大，说明该任务如果不及时插入，后续可选方案变差的风险越高，因此系统优先插入遗憾值较大的任务。")
    add_paragraph(doc, "步骤54：执行破坏操作。破坏算子包括随机移除、最差成本移除和高能耗移除。随机移除用于扩大搜索空间；最差成本移除用于重新优化当前成本较高的任务；高能耗移除用于优先降低无人机与无人车能耗。")
    add_paragraph(doc, "步骤55：执行修复操作。修复算子包括贪婪插入、遗憾值插入和接力感知遗憾值插入。接力感知遗憾值插入在比较成本增量的同时考虑模式风险、预计等待时间和预计时间裕度。")
    add_paragraph(doc, "步骤56：采用模拟退火准则接受新解。本实验项目中ALNS默认最大迭代次数为30，初始温度为100，降温率为0.95。若新解成本低于当前解，则直接接受；若新解成本高于当前解，则按以下概率接受：")
    add_acceptance_equation(doc)
    add_paragraph(doc, "其中，Pacc表示接受概率，Fold表示旧解成本，Fnew表示新解成本，τ表示当前温度。该机制使算法在前期保持一定探索能力，在后期逐步收敛。")
    add_paragraph(doc, "步骤57：更新算子权重。系统根据破坏算子和修复算子的历史表现自适应更新算子权重。若产生历史最优解，则对应算子权重增加0.5；若仅改善当前解，则权重增加0.1；若未产生改进，则权重乘以0.95衰减，并保持不低于0.1。其一般形式表示为：")
    add_weight_update_equation(doc)
    add_paragraph(doc, "其中，Q表示算子权重，ρ表示更新系数，S表示当前统计周期内算子的表现评分。通过该机制，ALNS逐步偏向更适合当前任务分布和设备状态的算子组合。")

    add_paragraph(doc, "具体的，所述步骤6具体包括以下子步骤：")
    add_paragraph(doc, "步骤61：检测等待无人车状态任务。系统首先遍历状态为waiting_for_agv的任务，该类任务已经选择接力配送，但尚未完成无人机与无人车的接力衔接。")
    add_paragraph(doc, "步骤62：检测回退触发条件。本实验项目中的回退触发条件包括无人机电量低于20、等待时间超过10、中继点位置无效、任务剩余时间过短、部署方案不可行、预计等待时间过长以及预计时间裕度不足。触发条件可概括表示为：")
    add_fallback_trigger_equation(doc)
    add_paragraph(doc, "步骤63：执行中继点重选。触发回退后，系统重新遍历当前可用无人车及其候选中继点，并复用统一评价器计算成本增量、模式风险、预计等待时间和预计时间裕度。最优重选中继点表示如下：")
    add_reselect_equation(doc)
    add_paragraph(doc, "若存在新的可行接力方案，则任务保持waiting_for_agv状态，并更新中继点和无人车；若不存在新的可行接力方案，则系统评价直接配送方案，直接配送可行时将任务切换为in_progress状态；若接力方案和直接配送方案均不可行，则将任务状态恢复为pending，等待后续调度周期重新处理。")

    add_paragraph(doc, "具体的，所述步骤7具体包括以下子步骤：")
    add_paragraph(doc, "步骤71：输出任务分配结果。系统输出每个任务对应的无人机编号、任务编号、配送模式、无人车编号、中继点坐标和候选方案成本。若配送模式为直接配送，则无人车编号和中继点坐标为空；若配送模式为接力配送，则同步输出无人车移动至中继点的动作。")
    add_paragraph(doc, "步骤72：输出调度事件。对于接力配送任务，系统输出RELAY_REQUEST事件；对于执行阶段重新选择中继点的任务，系统输出RELAY_RESELECTED事件；对于由接力配送切换为直接配送的任务，系统输出RELAY_TO_DIRECT_FALLBACK事件；对于当前不可行任务，系统输出INFEASIBLE_TASK事件。")
    add_paragraph(doc, "步骤73：输出统计结果。系统统计直接配送数量、接力配送数量、回退次数、重新规划次数、不可行任务数量和最终方案总成本。上述统计结果可以用于实验对比、消融分析和后续参数调整。")
    add_paragraph(doc, "步骤74：生成实验复核信息。系统保留候选池构建、候选方案排序、算子权重变化、最优解成本变化和回退原因，使实验结果能够被复核。通过上述输出，本实施例不仅给出最终调度方案，还给出方案形成过程中的关键判断依据。")
    add_paragraph(doc, "本实施例通过步骤1至步骤7，将任务建模、候选中继点生成、可行性筛选、统一评价、候选池构建、ALNS联合优化和执行阶段回退连接为一个完整流程。该流程以本实验项目中的实际代码结构为基础，能够在同一评价口径下比较直接配送与接力配送，并在任务执行过程中根据设备状态变化进行可解释的回退处理。")


def build_document() -> Document:
    doc = Document()
    normalize_doc(doc)

    add_title(doc, "一种基于ALNS的无人机无人车协同绿色配送优化方法")

    add_heading(doc, "技术领域")
    add_paragraph(
        doc,
        "本发明涉及低空物流配送、无人机无人车协同调度、绿色路径优化和智能物流决策技术领域，尤其涉及一种基于自适应大邻域搜索算法的无人机无人车协同绿色配送优化方法。该方法面向园区物流、校园配送、厂区物料转运、封闭或半封闭区域末端配送等应用场景，通过统一建模任务、无人机、无人车、候选中继点、能耗成本、等待成本和回退风险，实现直接配送方式与接力配送方式的统一评价和联合优化。"
    )

    add_heading(doc, "背景技术")
    for text in [
        "随着即时配送、智慧园区、低空经济和无人配送技术的发展，末端配送系统逐渐从单一地面车辆配送向空地协同配送演进。无人车具有续航时间长、载重能力强、运行稳定和适合执行地面主干运输任务等特点，但在道路绕行、区域封闭、局部拥堵或路径受限的情况下，容易出现任务响应慢、路径冗长和单位时间配送效率不足的问题。无人机具有飞行路径灵活、跨越障碍能力强和点到点响应速度快等优势，适合执行高时效、小批量和短距离配送任务，但无人机受到电池容量、载荷能力、飞行安全边界、禁飞区和起降条件等因素限制，难以独立承担全部配送任务。",
        "在无人机无人车协同配送场景中，同一配送任务可以由无人机直接完成，也可以由无人车与无人机通过接力方式共同完成。接力配送通常由无人车靠近任务走廊或中继区域，无人机从当前位置飞往中继点，再由无人机完成取货、投送和返回衔接。该模式能够在一定条件下减少无人机长距离空飞，提高无人车地面运输资源利用率，并降低整体配送能耗。然而，接力配送不是简单地选择一个几何距离较短的中继点，还需要同时考虑无人机当前位置、无人车当前位置、任务起点、任务终点、双方到达中继点的时间差、无人机剩余航程、任务截止时间和后续任务插入位置。",
        "现有技术中，一类方法通常预设固定中继点或候选起降点，再根据距离或局部时间成本选择配送模式。该类方法实现简单，但容易忽略无人机从实际当前位置部署到中继点的迁移成本，导致接力方案评分偏低。另一类方法仅在任务分配层面粗略地区分无人机任务和无人车任务，没有将接力候选点生成、路径插入、等待同步、能耗约束和异常回退纳入统一模型。还有一些启发式方法虽然可以求解多任务调度问题，但对不同配送模式采用不同评价口径，导致直接配送和接力配送之间缺乏稳定、可解释的比较基础。",
        "在实际部署中，空地协同配送还会受到电量下降、无人车延迟、中继点失效、任务时限变化和环境约束变化等因素影响。若系统缺少统一的回退机制，则一旦接力方案不可执行，调度系统往往只能临时切换为直接配送或重新等待人工干预，容易造成任务延误、能耗增加和调度结果不可追溯。因此，需要一种能够贯穿候选生成、可行性筛选、统一评价、ALNS联合优化和执行阶段回退的协同配送优化方法，使算法评分结果与实际执行过程保持一致。",
        "本项目的实际实现已经包含任务对象、无人机对象、无人车对象、环境有效性判断、候选中继点生成、统一成本评价器、ALNS候选池、破坏算子、修复算子、自适应算子权重、模拟退火接受准则以及执行阶段回退逻辑。本发明在上述实现基础上，将直接配送和接力配送统一到同一候选空间中，以综合成本为核心评价指标，兼顾时间、能耗、碳排放估计、等待同步和回退风险，从而提高协同配送方案的绿色性、稳定性和工程可实施性。"
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "发明内容")
    for text in [
        "本发明的目的在于克服现有无人机无人车协同配送方法中接力点选择静态、配送模式评价口径不一致、无人机部署成本容易被低估、异常回退缺乏统一依据等不足，提供一种基于ALNS的无人机无人车协同绿色配送优化方法。该方法以实际项目中的取送货任务、无人机、无人车和配送环境为基础，统一生成直接配送候选方案和无人机无人车接力配送候选方案，并通过同一评价器完成候选筛选、路径插入、ALNS修复和执行阶段回退判断。",
        "相较于仅按照静态中继点或单一距离指标进行任务分配的方法，本发明至少具有以下改进：接力候选点由无人车当前位置、任务起点、任务终点、任务走廊采样点以及无人车到任务走廊的投影点共同生成，并与具体无人车绑定；接力评价中显式计入无人机从当前位置飞往中继点的部署段，避免默认无人机已经位于中继点；统一评价模型同时考虑时间成本、无人机能耗、无人车能耗、碳排放估计、等待惩罚和回退风险；ALNS求解中采用候选池、破坏算子、修复算子和自适应算子权重，提高联合搜索效率；执行阶段复用统一评价器进行中继点重选或直接配送回退。",
        "本发明的目的是通过以下技术方案来实现的：一种基于ALNS的无人机无人车协同绿色配送优化方法，包括以下步骤。",
        "步骤1：获取配送任务集合、无人机集合、无人车集合以及配送环境约束信息，建立空地协同配送基础数据模型。",
        "步骤2：针对任务与无人机组合构建候选服务方式集合，所述候选服务方式包括无人机直接配送方式和无人机无人车接力配送方式。",
        "步骤3：对候选服务方式进行可行性筛选，剔除不满足环境有效性、剩余航程、任务时限和接力部署条件的候选方案。",
        "步骤4：构建统一评价模型，对直接配送候选方案和接力配送候选方案计算综合成本，并输出可行性标志、成本分解、预计等待时间、预计时间裕度和模式风险值。",
        "步骤5：基于统一评价模型构建候选池，并采用ALNS对任务分配、设备匹配、执行模式、路径插入位置和中继点选择进行联合优化。",
        "步骤6：在执行过程中监测接力方案状态，当出现电量不足、等待超时、中继点失效、部署不可行、同步偏差过大或时间裕度不足时，触发中继点重选或直接配送回退。",
        "步骤7：输出无人机无人车协同配送方案，所述协同配送方案包括任务分配结果、执行模式、中继点位置、无人机执行顺序、无人车配合顺序、综合成本、能耗评估和时序评估结果。",
        "具体的，所述步骤2中，接力候选点至少包括无人车当前位置、任务起点、任务终点、任务走廊上的多个采样点以及无人车当前位置到任务走廊线段的投影点。候选点生成后按照坐标近似值进行去重，并保留与无人车绑定的可行候选点。",
        "具体的，所述步骤3中，对接力候选方案至少执行四类判断：判断候选中继点是否位于有效配送区域内；判断无人机从当前位置飞往候选中继点所需航程是否小于剩余可用航程；判断无人机到达中继点后是否仍具有完成取货、投送和返回衔接的航程余量；当任务设置截止时间时，判断候选方案预计完成时间是否满足任务时限要求。",
        "具体的，所述步骤4中，直接配送候选方案计算无人机飞行时间、无人机能耗、碳排放估计和基础回退风险；接力配送候选方案计算无人机部署段和配送段的飞行时间与能耗、无人车行驶能耗、碳排放估计、双方到达中继点的预计等待时间以及回退风险。模式风险值根据剩余航程裕度、任务时间裕度和无人机无人车同步偏差共同确定。",
        "具体的，所述步骤5中，ALNS使用的破坏算子至少包括随机移除、最差成本移除和高能耗移除，修复算子至少包括贪婪插入、遗憾值插入和接力感知遗憾值插入。算法根据破坏算子和修复算子的历史表现自适应调整算子权重，并采用模拟退火接受准则判断是否接受新解。",
        "具体的，所述步骤6中，对于处于等待无人车状态的任务，系统检测无人机电量、等待时间、候选中继点有效性和任务剩余时间。当触发异常条件时，系统重新遍历当前可用无人车及其候选中继点，并复用统一评价模型计算各接力候选方案；若存在新的可行接力方案，则选择综合评价最优的中继点和无人车执行重选；若不存在可行接力方案，则评价直接配送方案并在可行时切换为直接配送。",
        "本发明相对于现有技术具有如下有益效果：将直接配送和接力配送纳入同一候选空间和同一评价模型，避免不同执行模式之间评价口径不一致；在接力方案中显式考虑无人机从当前位置到中继点的部署段，使评分结果更接近真实执行过程；通过候选池和多类ALNS算子降低搜索复杂度，并提升任务分配、设备匹配和中继点选择的联合优化效果；通过等待惩罚、时间裕度和回退风险提高异常情况下方案切换的可解释性；综合考虑无人机能耗、无人车能耗和碳排放估计，有利于形成更符合绿色配送目标的协同调度方案。"
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "附图说明")
    for text in [
        "图一为本发明一种基于ALNS的无人机无人车协同绿色配送优化方法的总体流程示意图。",
        "图二为本发明中直接配送方式与接力配送方式的统一候选评估示意图。",
        "图三为本发明中接力配送方式的连续部署执行时序示意图。",
        "图四为本发明中ALNS破坏与修复迭代优化过程示意图。",
        "图五为本发明中执行阶段回退与中继点重选流程示意图。"
    ]:
        add_paragraph(doc, text)

    add_implementation_section(doc)
    return doc

    add_heading(doc, "具体实施方法")
    for text in [
        "为使本发明的目的、技术方案和有益效果更加清楚，以下结合实施例对本发明作进一步说明。所描述的实施例用于解释本发明的技术方案，不用于限定本发明的保护范围。本领域技术人员在不脱离本发明核心思想的前提下，对任务类型、设备参数、候选点数量、成本权重或算子配置进行等同替换，均应视为落入本发明的保护范围。",
        "实施例一：本实施例以园区末端配送场景为例。配送环境包括可通行区域、障碍区域、禁飞区域、任务服务点、无人机集合和无人车集合。任务为取送货任务，每个任务包括任务起点、任务终点、载荷、优先级和截止时间。无人机包括当前位置、电量、最大航程和最大飞行速度。无人车包括当前位置、行驶状态和最大行驶速度。系统以调度周期为单位读取当前环境状态，生成候选方案，执行统一评价，并通过ALNS获得当前周期的协同配送方案。",
        "在本实施例中，配送任务集合、无人机集合和无人车集合分别由项目中的任务对象、无人机对象和无人车对象表示。环境对象提供位置有效性判断，用于检测候选点是否位于可执行区域内。若候选点落入障碍区域、禁飞区域或环境边界外，则该候选点不参与后续评价。通过上述对象化建模，调度算法能够在每个周期读取真实设备状态，而不是基于固定仓库点或静态假设进行估算。",
        "第一部分为场景建模与基础对象定义。系统首先采集任务起点、任务终点、载荷、优先级、截止时间、无人机当前位置、无人机剩余电量、无人机最大航程、无人机飞行速度、无人车当前位置、无人车行驶状态以及环境边界、障碍区和禁飞区信息。然后将上述信息组织为任务集合、无人机集合、无人车集合和环境约束集合。对于无效坐标、不可执行任务、不可用设备和不满足基本约束的数据，系统进行预处理，形成候选生成和统一评价可直接调用的输入。",
        "第二部分为候选服务方式生成。对于每一个任务与无人机组合，系统生成直接配送候选方案。直接配送候选方案以无人机当前位置作为真实起点，依次经过任务起点和任务终点，并返回当前锚点或后续衔接位置。该路径语义用于避免将仓库点错误地作为无人机实际起点，从而减少直接配送成本估计偏差。",
        "对于接力配送候选方案，系统针对每辆空闲无人车生成候选中继点。候选中继点来源包括无人车当前位置、任务起点、任务终点、任务起点与任务终点之间任务走廊的多个采样点，以及无人车当前位置到任务走廊线段的投影点。任务走廊采样点用于覆盖任务路径中段接力情形，投影点用于获得无人车靠近任务走廊的接力位置，无人车当前位置用于保留近端接力方案，任务起点和任务终点用于覆盖取货点或投送点附近接力方案。",
        "候选中继点生成后，系统按照坐标近似值进行去重，避免多个来源生成同一或近似中继点造成重复评价。候选中继点与具体无人车绑定，形成中继点和无人车组合。该绑定关系用于后续同时计算无人车到达中继点的地面行驶成本和无人机到达中继点的空中部署成本。若当前没有空闲无人车，则系统仅保留直接配送候选方案。",
        "第三部分为可行性筛选。系统对接力候选方案执行多级判断。首先判断候选中继点是否为环境有效位置。其次计算无人机从当前位置飞往候选中继点的部署距离，并根据无人机剩余电量和最大航程得到剩余可用航程。如果部署距离已经超过剩余可用航程，则候选方案被剔除。再次计算无人机到达中继点后完成中继点到任务起点、任务起点到任务终点、任务终点返回中继点的配送段距离，如果剩余航程不足以覆盖该配送段，则候选方案被剔除。最后，当任务具有截止时间时，系统估算候选方案完成时间，若超过截止时间，则候选方案不进入统一评价阶段。",
        "在本实施例中，接力配送总飞行距离由无人机部署段和无人机配送段组成。部署段表示无人机从当前位置飞往候选中继点的距离，配送段表示无人机从中继点飞往任务起点、任务终点并返回中继点的距离。其表达式为："
    ]:
        add_paragraph(doc, text)
    add_distance_equation(doc, "relay", [("U", "R"), ("R", "S"), ("S", "E"), ("E", "R")])
    add_paragraph(doc, "其中，U表示无人机当前位置，R表示候选中继点，S表示任务起点，E表示任务终点，D表示两个位置之间的欧氏距离。该公式用于说明接力配送评价必须包含无人机从当前位置飞往中继点的部署段，而不是只计算中继点之后的配送段。")
    add_paragraph(doc, "直接配送总飞行距离由无人机当前位置到任务起点、任务起点到任务终点、任务终点返回当前锚点或后续衔接位置构成。其表达式为：")
    add_distance_equation(doc, "direct", [("U", "S"), ("S", "E"), ("E", "H")])
    add_paragraph(doc, "其中，H表示返回锚点或后续衔接位置。当系统不要求返回固定点时，H可以取为任务终点或下一任务衔接点。通过上述距离定义，直接配送和接力配送均以真实执行路径作为成本计算基础。")

    for text in [
        "第四部分为统一评价模型。统一评价模型的输入包括候选服务方式、无人机、任务、无人车、候选中继点、当前无人机路线、当前无人车路线和插入位置。统一评价模型的输出包括综合成本、可行性标志、成本分解、预计等待时间、预计时间裕度和模式风险值。该评价器同时服务于候选池构建、初始解生成、ALNS修复插入和执行阶段回退判断。",
        "对于直接配送方案，统一评价模型计算无人机飞行时间、无人机能耗、碳排放估计和基础回退风险。对于接力配送方案，统一评价模型计算无人机部署段和配送段的飞行时间与能耗、无人车到中继点的行驶能耗、碳排放估计、无人机和无人车到达中继点的预计等待时间以及由航程裕度、时间裕度和同步偏差共同形成的模式风险。",
        "在本实施例中，无人机能耗和无人车能耗采用距离与单位距离能耗参数相乘的简化模型。该模型与项目中的能耗评价接口保持一致，可以在未接入更复杂动力学模型时提供稳定、可解释的成本估算。能耗表达式为："
    ]:
        add_paragraph(doc, text)
    add_energy_equation(doc, "uav", "uav", "uav")
    add_energy_equation(doc, "agv", "agv", "agv")
    add_paragraph(doc, "其中，Euav表示无人机能耗，Eagv表示无人车能耗，Kuav表示无人机单位距离能耗参数，Kagv表示无人车单位距离能耗参数，Luav表示无人机飞行距离，Lagv表示无人车行驶距离。上述公式为 Word 原生公式对象，用于表征绿色配送目标中的能耗成本。")
    add_paragraph(doc, "等待时间用于衡量无人机与无人车在候选中继点的同步偏差。若无人机先到达中继点，则无人机需要等待无人车；若无人车先到达中继点，则无人车需要等待无人机。等待时间表达式为：")
    add_wait_equation(doc)
    add_paragraph(doc, "其中，Tuav表示无人机预计到达中继点的时间，Tagv表示无人车预计到达中继点的时间。等待时间越大，说明接力同步性越差，接力方案发生延误或回退的风险越高。")
    add_paragraph(doc, "模式风险由航程裕度、时间裕度和同步偏差共同确定。剩余航程越接近所需航程，风险越高；任务剩余时间越少，风险越高；等待时间越长，风险越高。模式风险表达式为：")
    add_risk_equation(doc)
    add_paragraph(doc, "其中，A表示航程裕度风险项，B表示时间裕度风险项，C表示同步偏差风险项。上述权重可根据具体应用场景调整，但在本实施例中使用固定权重以保持评价过程透明。综合成本表达式为：")
    add_cost_equation(doc)
    add_paragraph(doc, "其中，F表示候选方案综合成本，T表示时间成本，Ccarbon表示碳排放估计，W表示等待惩罚，R表示模式风险，wt、wu、wa、wc、ww和wr分别表示各成本项权重。通过上述综合成本，系统可以在同一评价口径下比较直接配送方案和接力配送方案。")

    for text in [
        "第五部分为候选池构建。系统针对每个任务与无人机组合建立候选池。候选池包括直接配送可行标志和若干接力候选方案。接力候选方案首先由候选中继点生成器生成，再由统一评价模型计算成本、风险和等待时间。系统对可行候选方案按照综合成本、模式风险和预计等待时间进行排序，并选取前若干个候选方案进入ALNS求解阶段。",
        "为了避免候选方案全部集中在同一无人车或同一局部区域，本实施例在选取候选池时优先保持无人车来源多样性。具体而言，系统先从排序后的候选方案中选择不同无人车对应的优质候选点；若数量不足，再从剩余候选方案中按综合评价顺序补足。该处理能够在候选池规模受限的情况下保留更多潜在可行协同方式。",
        "第六部分为ALNS联合优化。系统首先采用遗憾值插入方式生成初始解。对于每个未分配任务，系统收集其在不同无人机、不同配送模式和不同中继点下的可行候选方案，并按照综合成本排序。若某任务的最优方案与次优方案成本差距较大，说明该任务对插入位置和服务模式较敏感，系统优先插入该任务，以降低后续难以分配的风险。",
        "在ALNS迭代过程中，系统通过破坏算子从当前解中移除部分任务，再通过修复算子重新插入任务。破坏算子包括随机移除、最差成本移除和高能耗移除。随机移除用于扩大搜索空间；最差成本移除用于优先处理当前解中成本较高的分配；高能耗移除用于优先重新优化能耗较高的任务分配。修复算子包括贪婪插入、遗憾值插入和接力感知遗憾值插入。接力感知遗憾值插入在比较成本差异的同时，还考虑模式风险、预计等待时间和预计时间裕度。",
        "系统根据算子历史表现自适应调整算子权重。当某一组破坏算子和修复算子产生历史最优解时，系统提高其权重；当仅改善当前解时，系统小幅提高其权重；当未产生改进时，系统适度衰减其权重。通过该机制，ALNS能够逐步偏向更适合当前场景的算子组合。每轮迭代产生新解后，系统采用模拟退火接受准则判断是否接受该解。当新解成本更低时直接接受；当新解成本更高时，根据温度参数以一定概率接受，以避免搜索陷入局部最优。温度随迭代次数逐步降低，算法由探索阶段过渡到收敛阶段。",
        "第七部分为执行阶段回退机制。对于已经选择接力配送的任务，系统不会假定接力方案一定成功执行，而是在执行阶段持续监测无人机电量、等待时间、候选中继点有效性和任务剩余时间。当无人机电量低于阈值、等待时间超过阈值、中继点失效、部署不可行、同步偏差过大或时间裕度不足时，系统触发回退判断。",
        "触发回退后，系统首先遍历当前可用无人车及其候选中继点，复用统一评价模型计算各接力候选方案的综合成本、模式风险、预计等待时间和预计时间裕度。如果存在新的可行接力方案，系统选择综合评价最优的候选中继点和无人车执行重选。如果不存在可行接力方案，则系统评价直接配送方案，并在直接配送可行时将任务切换为直接配送。如果接力方案和直接配送方案均不可行，则系统将任务保持为待分配状态，等待后续调度周期重新处理。",
        "第八部分为输出结果。系统输出的协同配送方案包括任务与无人机的匹配关系、任务与无人车的配合关系、配送模式、中继点位置、无人机执行顺序、无人车移动动作、综合成本、成本分解、预计等待时间、模式风险和回退事件。上述结果可以用于调度执行，也可以用于后续分析不同策略的任务完成效率、能耗水平和回退频率。",
        "在一个具体示例中，任务起点位于园区仓储点，任务终点位于用户服务点，无人机当前位置不一定与仓库点重合，无人车当前位置位于园区道路网络中。系统先根据任务起点和任务终点形成任务走廊，再生成中继候选点。若某候选点使无人机部署距离较短、无人车行驶距离适中、双方到达时间差较小，并且无人机剩余航程能够覆盖部署段和配送段，则该候选点进入候选池。随后ALNS在直接配送和多个接力配送候选方案之间搜索，最终输出当前综合成本最低且满足约束条件的方案。",
        "本实施例中的距离模型、能耗模型、成本权重、候选池规模、最大迭代次数、初始温度、降温率、破坏数量和修复数量均可以根据实际场景进行调整。上述参数调整不改变本发明的核心思想，即以真实当前位置为起点生成直接配送和接力配送候选方案，通过统一评价器比较不同执行模式，并利用ALNS完成任务分配、设备匹配、路径插入和中继点选择的联合优化。",
        "综上，本发明通过候选中继点生成、可行性筛选、统一评价、候选池构建、ALNS联合优化和执行阶段回退机制，形成了完整的无人机无人车协同绿色配送优化流程。该流程能够减少接力方案评分与执行之间的偏差，提高不同配送模式之间的可比性，降低搜索复杂度，并在异常情况下提供可解释的回退依据，适用于园区物流、校园配送、厂区配送和其他多约束末端配送场景。"
    ]:
        add_paragraph(doc, text)

    for text in [
        "本发明通过统一候选空间使直接配送和接力配送能够在同一评价口径下比较，避免因评价模型不同导致模式选择不稳定。",
        "本发明在接力方案中显式计入无人机从当前位置到中继点的部署距离、部署时间和部署能耗，使接力评分更接近实际执行过程。",
        "本发明通过候选池筛选和多类ALNS算子降低联合搜索复杂度，并通过自适应算子权重提高搜索稳定性。",
        "本发明通过等待时间、时间裕度和模式风险刻画接力同步性和回退风险，使异常情况下的中继点重选或直接配送切换具有明确依据。",
        "本发明综合考虑无人机能耗、无人车能耗和碳排放估计，有利于在满足配送时效的同时降低单位任务能耗，提升空地协同配送的绿色性和工程应用价值。"
    ]:
        add_paragraph(doc, text)

    return doc


def count_chinese_chars(path: str) -> int:
    doc = Document(path)
    text = "".join(p.text for p in doc.paragraphs)
    return len(re.findall(r"[\u4e00-\u9fff]", text))


def validate(path: str) -> dict:
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    return {
        "omath_count": xml.count("<m:oMath") + xml.count("<m:oMathPara"),
        "has_raw_latex_backslash": "\\" in xml,
        "has_bracket_paragraph_numbers": bool(re.search(r"\[\d{4}\]", xml)),
        "chinese_chars": count_chinese_chars(path),
    }


def main() -> None:
    doc = build_document()
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(validate(OUTPUT_PATH))


if __name__ == "__main__":
    main()
