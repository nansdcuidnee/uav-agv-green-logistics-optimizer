from __future__ import annotations

import os
import zipfile
from collections import Counter

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


DOCX_PATH = r"D:\我的专利（1）\ALNS说明书 .docx"


def text_of(paragraph: Paragraph) -> str:
    return paragraph.text.strip().replace("\u3000", " ")


def clear_and_set(paragraph: Paragraph, text: str) -> None:
    element = paragraph._element
    for child in list(element):
        element.remove(child)
    if text:
        paragraph.add_run(text)


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if paragraph.style is not None:
        new_para.style = paragraph.style
    new_para.add_run(text)
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def find_exact(paragraphs: list[Paragraph], expected: str) -> int:
    for idx, paragraph in enumerate(paragraphs):
        if text_of(paragraph) == expected:
            return idx
    raise ValueError(f"Paragraph not found: {expected}")


def find_startswith(paragraphs: list[Paragraph], prefix: str) -> int:
    for idx, paragraph in enumerate(paragraphs):
        if text_of(paragraph).startswith(prefix):
            return idx
    raise ValueError(f"Paragraph with prefix not found: {prefix}")


def simplify_invention_content(doc: Document) -> None:
    paragraphs = doc.paragraphs
    start = find_exact(paragraphs, "在步骤一中，具体包括以下三个步骤：")
    end = find_startswith(paragraphs, "[0015] 进一步地")
    for paragraph in list(paragraphs[start:end]):
        remove_paragraph(paragraph)


def insert_detail_after_heading(doc: Document, heading: str, marker: str, lines: list[str]) -> None:
    paragraphs = doc.paragraphs
    if any(marker in text_of(paragraph) for paragraph in paragraphs):
        return
    heading_idx = find_exact(paragraphs, heading)
    anchor = paragraphs[heading_idx]
    for line in reversed(lines):
        insert_after(anchor, line)


def rebuild_implementation_detail(doc: Document) -> None:
    insert_detail_after_heading(
        doc,
        "第一部分为：场景建模与基础对象定义",
        "步骤11，采集任务起点",
        [
            "在第一部分中，具体包括以下三个步骤：",
            "步骤11，采集任务起点、任务终点、载荷、优先级、截止时间、无人机当前位置、无人机剩余电量、无人车当前位置以及配送环境约束信息。",
            "步骤12，依据采集到的任务信息、设备信息和环境信息，建立任务集合、无人机集合、无人车集合以及环境约束集合；其中，任务集合中的每个任务至少包含任务起点、任务终点、载荷、优先级和截止时间，无人机集合中的每架无人机至少包含当前位置、电量、最大航程和飞行速度，无人车集合中的每辆无人车至少包含当前位置、行驶状态和行驶速度。",
            "步骤13，对缺失数据、越界坐标、不可执行设备和不满足基本服务条件的任务进行预处理，形成后续候选生成阶段可直接调用的标准化输入。",
        ],
    )

    insert_detail_after_heading(
        doc,
        "第二部分为：候选服务方式生成与可行性筛选",
        "步骤21，为每一个任务与无人机组合生成无人机直接配送候选方案",
        [
            "在第二部分中，候选服务方式生成具体包括以下三个步骤：",
            "步骤21，为每一个任务与无人机组合生成无人机直接配送候选方案，所述直接配送候选方案以无人机当前位置作为起点，并以任务起点、任务终点以及返回锚点或后续衔接点构成完整服务路径。",
            "步骤22，基于无人车当前位置、任务起点、任务终点、任务走廊采样点以及无人车到任务走廊的投影点生成接力配送候选中继点。",
            "步骤23，对重复中继点、落入禁飞区或障碍区域的中继点以及不满足接力部署条件的中继点进行剔除，得到候选服务方式集合。",
            "在第二部分中，可行性筛选具体包括以下四个步骤：",
            "步骤24，判断候选中继点是否位于有效可执行区域内，若否，则剔除该候选方案。",
            "步骤25，判断无人机是否能够从当前位置连续飞行至候选中继点，若否，则剔除该候选方案。",
            "步骤26，判断无人机到达候选中继点后的剩余航程是否足以完成配送段，若否，则剔除该候选方案。",
            "步骤27，判断候选方案的预计完成时间是否满足任务截止时间要求，若满足，则将该候选方案输入统一评价模型。",
        ],
    )

    insert_detail_after_heading(
        doc,
        "第四部分为：基于ALNS的联合优化过程",
        "步骤41，按照综合增量成本由低到高的顺序构造初始解",
        [
            "在第四部分中，基于ALNS的联合优化具体包括以下五个步骤：",
            "步骤41，按照综合增量成本由低到高的顺序构造初始解。具体地，对每个未分配任务遍历其可行候选服务方式，并分别计算插入至各无人机路线和各无人车配合序列后的综合增量成本、等待变化量以及风险变化量。",
            "步骤42，通过随机移除、最差成本移除或风险导向移除执行破坏操作。其中，随机移除用于提高解空间探索能力，最差成本移除用于优先剔除导致综合成本偏高的任务分配，风险导向移除用于优先处理续航边际不足、等待偏差过大或回退概率较高的接力任务。",
            "步骤43，通过贪婪插入、遗憾值插入或接力感知修复执行恢复操作。具体地，对每个待恢复任务重新枚举直接配送方式和接力配送方式下的可行插入位置，并基于统一评价器重新计算插入后的综合增量成本。",
            "步骤44，结合模拟退火接受准则判断是否接受新解，并更新当前解和历史最优解。若新解综合成本优于当前解，则直接接受；若新解劣于当前解，则按照与温度相关的概率接受。",
            "步骤45，当达到最大迭代次数、连续若干轮迭代未能改善当前最优解、温度参数降低至预设下限，或者当前解与最优解之间的改进幅度小于预设阈值时，输出历史最优解作为最优或近优联合优化结果。",
        ],
    )


def remove_remaining_duplicate_text(doc: Document) -> int:
    texts = [text_of(p) for p in doc.paragraphs]
    counts = Counter(t for t in texts if t)
    seen: set[str] = set()
    removed = 0
    for paragraph in list(doc.paragraphs):
        current = text_of(paragraph)
        if not current or counts[current] <= 1:
            continue
        if current in seen and not current.startswith("["):
            remove_paragraph(paragraph)
            removed += 1
        else:
            seen.add(current)
    return removed


def validate(path: str) -> dict[str, int | bool]:
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    return {
        "omath_count": xml.count("<m:oMath") + xml.count("<m:oMathPara"),
        "has_raw_latex": "\\" in xml,
        "paragraph_count": xml.count("<w:p"),
        "has_invention_step11": "在步骤一中，具体包括以下三个步骤" in xml,
    }


def main() -> None:
    if not os.path.exists(DOCX_PATH):
        raise FileNotFoundError(DOCX_PATH)

    before = validate(DOCX_PATH)
    doc = Document(DOCX_PATH)
    simplify_invention_content(doc)
    rebuild_implementation_detail(doc)
    removed = remove_remaining_duplicate_text(doc)
    doc.save(DOCX_PATH)
    after = validate(DOCX_PATH)
    print({"before": before, "after": after, "removed_duplicates": removed})


if __name__ == "__main__":
    main()
