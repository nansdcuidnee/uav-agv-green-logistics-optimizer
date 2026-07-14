from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


TARGET_DOC = r"D:\uav-agv-green-logistics-optimizer\docs\权利要求书_一种基于ALNS的无人机无人车协同绿色配送优化方法.docx"
OUTPUT_DOC = r"D:\uav-agv-green-logistics-optimizer\docs\权利要求书_一种基于ALNS的无人机无人车协同绿色配送优化方法.new.docx"


def clear_document(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_run_font(run, size=12, bold=False) -> None:
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        rpr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "宋体")


def add_text_paragraph(doc: Document, text: str, center=False, bold=False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, 12, bold)


def mr(text: str):
    run = OxmlElement("m:r")
    text_node = OxmlElement("m:t")
    text_node.text = text
    run.append(text_node)
    return run


def sub(base: str, subscript: str):
    node = OxmlElement("m:sSub")
    e = OxmlElement("m:e")
    e.append(mr(base))
    s = OxmlElement("m:sub")
    s.append(mr(subscript))
    node.append(e)
    node.append(s)
    return node


def sup(base: str, superscript: str):
    node = OxmlElement("m:sSup")
    e = OxmlElement("m:e")
    e.append(mr(base))
    s = OxmlElement("m:sup")
    s.append(mr(superscript))
    node.append(e)
    node.append(s)
    return node


def subsup(base: str, subscript: str, superscript: str):
    node = OxmlElement("m:sSubSup")
    e = OxmlElement("m:e")
    e.append(mr(base))
    sub_node = OxmlElement("m:sub")
    sub_node.append(mr(subscript))
    sup_node = OxmlElement("m:sup")
    sup_node.append(mr(superscript))
    node.append(e)
    node.append(sub_node)
    node.append(sup_node)
    return node


def math(nodes):
    omath = OxmlElement("m:oMath")
    for node in nodes:
        omath.append(node)
    return omath


def add_math_paragraph(doc: Document, nodes) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    para = OxmlElement("m:oMathPara")
    para.append(math(nodes))
    p._element.append(para)


def add_mixed_paragraph(doc: Document, parts) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    for kind, value in parts:
        if kind == "text":
            run = p.add_run(value)
            set_run_font(run)
        elif kind == "math":
            p._p.append(math(value))
        else:
            raise ValueError(kind)


def task_vector_nodes():
    return [
        sub("x", "i"), mr("=("), subsup("p", "i", "s"), mr(","),
        subsup("p", "i", "e"), mr(","), sub("q", "i"), mr(","),
        sub("π", "i"), mr(",["), sub("a", "i"), mr(","), sub("b", "i"), mr("])"),
    ]


def uav_state_nodes():
    return [sub("s", "u"), mr("=("), sub("p", "u"), mr(","), sub("B", "u"), mr(","), sub("Q", "u"), mr(","), sub("R", "u"), mr(")")]


def agv_state_nodes():
    return [sub("g", "v"), mr("=("), sub("p", "v"), mr(","), sub("B", "v"), mr(","), sub("Q", "v"), mr(")")]


def main() -> None:
    doc = Document(TARGET_DOC)
    clear_document(doc)
    if "Normal" in doc.styles:
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(12)
        doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    add_text_paragraph(doc, "权 利 要 求 书", center=True, bold=True)

    add_text_paragraph(
        doc,
        "1.一种基于ALNS的无人机无人车协同绿色配送优化方法，其特征在于，包括以下步骤：步骤1，获取配送任务信息、无人机与无人车设备状态信息以及配送环境约束信息；步骤2，基于所述配送任务信息、设备状态信息和环境约束信息，构建无人机直送方式和无人机与无人车接力配送方式的候选服务方式集合；步骤3，对所述候选服务方式集合进行可行性筛选，剔除不满足载重约束、续航约束、时效约束和安全约束的候选方案；步骤4，构建以总运营成本最小化为目标的综合评价模型，对通过可行性筛选的候选方案进行成本评估；步骤5，采用启发式插入方式生成初始解，并基于ALNS算法通过破坏算子与修复算子对所述初始解进行迭代优化，得到协同调度解；步骤6，当迭代优化过程或方案执行过程出现中继点失效、设备状态异常、任务动态变化或环境条件变化时，触发动态重规划与回退处理机制，对受影响任务重新进行方案评估、设备匹配和路径调整；步骤7，输出满足约束条件的无人机无人车协同绿色配送方案，所述配送方案至少包括任务分配结果、设备执行顺序、中继交接节点和配送执行时序。"
    )

    add_mixed_paragraph(doc, [
        ("text", "2.根据权利要求1所述的一种基于ALNS的无人机无人车协同绿色配送优化方法，其特征在于，所述步骤1具体包括以下子步骤：步骤11，构建待配送任务集合"),
        ("math", [mr("Ω={1,2,...,n}")]),
        ("text", "，对于任一配送任务"),
        ("math", [mr("i∈Ω")]),
        ("text", "，获取其任务属性向量"),
        ("math", task_vector_nodes()),
        ("text", "，其中，"),
        ("math", [subsup("p", "i", "s")]),
        ("text", "表示任务取货位置，"),
        ("math", [subsup("p", "i", "e")]),
        ("text", "表示任务送达位置，"),
        ("math", [sub("q", "i")]),
        ("text", "表示货物重量，"),
        ("math", [sub("π", "i")]),
        ("text", "表示任务优先级，"),
        ("math", [mr("["), sub("a", "i"), mr(","), sub("b", "i"), mr("]")]),
        ("text", "表示任务时间窗；步骤12，构建无人机集合"),
        ("math", [mr("U={1,2,...,m}")]),
        ("text", "，并获取任一无人机"),
        ("math", [mr("u∈U")]),
        ("text", "的状态向量"),
        ("math", uav_state_nodes()),
        ("text", "，其中，"),
        ("math", [sub("p", "u")]),
        ("text", "表示无人机当前位置，"),
        ("math", [sub("B", "u")]),
        ("text", "表示无人机剩余电量，"),
        ("math", [sub("Q", "u")]),
        ("text", "表示无人机额定载荷，"),
        ("math", [sub("R", "u")]),
        ("text", "表示无人机最大续航航程；步骤13，构建无人车集合"),
        ("math", [mr("V={1,2,...,l}")]),
        ("text", "，并获取任一无人车"),
        ("math", [mr("v∈V")]),
        ("text", "的状态向量"),
        ("math", agv_state_nodes()),
        ("text", "，其中，"),
        ("math", [sub("p", "v")]),
        ("text", "表示无人车当前位置，"),
        ("math", [sub("B", "v")]),
        ("text", "表示无人车剩余电量，"),
        ("math", [sub("Q", "v")]),
        ("text", "表示无人车载重能力；步骤14，构建环境约束集合"),
        ("math", [mr("E={"), sub("E", "road"), mr(","), sub("E", "forbid"), mr(","), sub("E", "fly"), mr(","), sub("E", "weather"), mr("}")]),
        ("text", "，其中，"),
        ("math", [sub("E", "road")]),
        ("text", "表示道路通行约束，"),
        ("math", [sub("E", "forbid")]),
        ("text", "表示地面禁行区域约束，"),
        ("math", [sub("E", "fly")]),
        ("text", "表示无人机禁飞区域约束，"),
        ("math", [sub("E", "weather")]),
        ("text", "表示气象约束，用于后续候选构建和约束筛选。"),
    ])

    add_mixed_paragraph(doc, [
        ("text", "3.根据权利要求2所述的一种基于ALNS的无人机无人车协同绿色配送优化方法，其特征在于，所述步骤2具体包括以下子步骤：步骤21，针对任务"),
        ("math", [mr("i")]),
        ("text", "和无人机"),
        ("math", [mr("u")]),
        ("text", "构建直送候选"),
        ("math", [subsup("h", "iu", "D"), mr("=(i,u)")]),
        ("text", "，所述直送候选表示无人机由当前位置或基地出发，依次完成取货、送达和返航；步骤22，针对任务"),
        ("math", [mr("i")]),
        ("text", "、无人机"),
        ("math", [mr("u")]),
        ("text", "、无人车"),
        ("math", [mr("v")]),
        ("text", "和中继节点"),
        ("math", [mr("r")]),
        ("text", "构建接力候选"),
        ("math", [subsup("h", "iuvr", "R"), mr("=(i,u,v,r)")]),
        ("text", "；步骤23，形成统一候选服务方式集合"),
        ("math", [sub("H", "iu"), mr("="), subsup("H", "iu", "D"), mr("∪"), subsup("H", "iu", "R")]),
        ("text", "，使直送候选和接力候选在同一评价模型中比较；步骤24，对每一候选服务方式记录其对应的任务编号、无人机编号、无人车编号、中继节点、执行顺序和预估成本，从而形成供ALNS算法直接调用的候选池。"),
    ])

    add_mixed_paragraph(doc, [
        ("text", "4.根据权利要求3所述的一种基于ALNS的无人机无人车协同绿色配送优化方法，其特征在于，所述接力候选中的中继节点由候选中继节点集合产生，所述候选中继节点集合为"),
        ("math", [sub("P", "vi"), mr("={"), sub("p", "v"), mr(","), subsup("p", "i", "s"), mr(","), subsup("p", "i", "e"), mr(","), sub("p", "mid"), mr(","), sub("p", "proj"), mr("}")]),
        ("text", "，其中，"),
        ("math", [sub("p", "v")]),
        ("text", "表示无人车当前位置，"),
        ("math", [subsup("p", "i", "s")]),
        ("text", "和"),
        ("math", [subsup("p", "i", "e")]),
        ("text", "分别表示取货位置和送达位置，"),
        ("math", [sub("p", "mid")]),
        ("text", "表示任务起终点连线上的采样点，"),
        ("math", [sub("p", "proj")]),
        ("text", "表示无人车位置到任务起终点连线的投影点；对所述候选中继节点集合进行去重、边界检查和禁飞禁行检查后，得到有效中继节点集合；并按照候选节点对应的预估行驶距离、预估飞行距离和等待风险进行排序，选取排名靠前的若干中继节点进入候选池，以减少后续搜索空间。"),
    ])

    add_text_paragraph(
        doc,
        "5.根据权利要求1所述的一种基于ALNS的无人机无人车协同绿色配送优化方法，其特征在于，所述步骤3具体包括以下子步骤：步骤31，对候选方案进行载重约束筛选；步骤32，对候选方案进行续航约束筛选；步骤33，对候选方案进行时间窗约束筛选；步骤34，对候选方案进行安全约束筛选；步骤35，对接力配送候选进行无人机与无人车同步约束筛选，以避免无人车未到达中继节点时无人机已进入起降或回收阶段。"
    )
    add_math_paragraph(doc, [
        sub("q", "i"), mr("≤"), sub("Q", "u"), mr("，"),
        sub("R", "req"), mr("≤"), sub("R", "u"), sub("B", "u"), mr("/100，"),
        sub("a", "i"), mr("≤"), sub("t", "i"), mr("≤"), sub("b", "i"), mr("，"),
        mr("r∉"), sub("E", "fly"), mr("∪"), sub("E", "forbid"),
    ])
    add_mixed_paragraph(doc, [
        ("text", "其中，"),
        ("math", [sub("R", "req")]),
        ("text", "表示无人机执行候选方案所需航程，"),
        ("math", [sub("t", "i")]),
        ("text", "表示任务预计完成时刻；当上述任一约束不满足时，将对应候选方案从候选服务方式集合中剔除；对于接力配送候选，还进一步比较无人车到达中继节点的时刻与无人机到达中继节点的时刻，当二者差值超过预设同步容忍阈值时，将该接力候选判定为不可行。"),
    ])

    add_text_paragraph(
        doc,
        "6.根据权利要求1所述的一种基于ALNS的无人机无人车协同绿色配送优化方法，其特征在于，所述步骤4中的综合评价模型以总运营成本最小化为目标，所述总运营成本由时间成本、无人机能耗、无人车能耗、碳排放成本、等待惩罚、超时惩罚、回退惩罚和未服务惩罚加权得到；其中，无人机能耗根据飞行距离、任务载荷和起降次数估算，无人车能耗根据地面行驶距离和载重状态估算，碳排放成本根据能耗结果与预设等效排放系数换算得到。"
    )
    add_math_paragraph(doc, [
        mr("min "), mr("C="), sub("w", "t"), mr("T+"), sub("w", "u"), sub("E", "u"), mr("+"),
        sub("w", "v"), sub("E", "v"), mr("+"), sub("w", "c"), sub("C", "c"), mr("+"),
        sub("w", "w"), sub("P", "w"), mr("+"), sub("w", "o"), sub("P", "o"), mr("+"),
        sub("w", "f"), sub("P", "f"), mr("+"), sub("w", "n"), sub("P", "n"),
    ])
    add_mixed_paragraph(doc, [
        ("text", "其中，"),
        ("math", [mr("T")]),
        ("text", "表示配送时长成本，"),
        ("math", [sub("E", "u")]),
        ("text", "和"),
        ("math", [sub("E", "v")]),
        ("text", "分别表示无人机能耗和无人车能耗，"),
        ("math", [sub("C", "c")]),
        ("text", "表示碳排放成本，"),
        ("math", [sub("P", "w")]),
        ("text", "、"),
        ("math", [sub("P", "o")]),
        ("text", "、"),
        ("math", [sub("P", "f")]),
        ("text", "和"),
        ("math", [sub("P", "n")]),
        ("text", "分别表示等待、超时、回退和未服务惩罚项，各"),
        ("math", [mr("w")]),
        ("text", "为对应权重。"),
    ])

    add_mixed_paragraph(doc, [
        ("text", "7.根据权利要求1所述的一种基于ALNS的无人机无人车协同绿色配送优化方法，其特征在于，所述步骤5中的初始解生成具体包括以下子步骤：步骤51，将待插入任务按照任务优先级和时间窗紧迫度排序；步骤52，针对每个待插入任务计算其在各候选服务方式下的插入增量成本"),
        ("math", [sub("ΔC", "i"), mr("(h)=C(S⊕h)-C(S)")]),
        ("text", "；步骤53，选择满足约束且插入增量成本最小的候选方案"),
        ("math", [sup("h", "*"), mr("=argmin "), sub("ΔC", "i"), mr("(h)")]),
        ("text", "加入当前解；步骤54，对无法插入的任务标记为待修复任务，并交由后续修复算子处理；步骤55，在插入过程中同步更新无人机飞行序列、无人车行驶序列、设备剩余电量和任务状态，使初始解能够作为ALNS迭代搜索的可行起点。"),
    ])

    add_text_paragraph(
        doc,
        "8.根据权利要求1所述的一种基于ALNS的无人机无人车协同绿色配送优化方法，其特征在于，所述ALNS算法中的破坏算子至少包括随机移除算子、最高成本移除算子和高能耗移除算子；其中，随机移除算子用于维持搜索多样性，最高成本移除算子用于移除对当前解成本贡献最大的任务或协同关系，高能耗移除算子用于优先移除造成无人机高能耗或碳排放增加的任务片段；每次破坏操作移除的任务数量根据当前解规模和预设破坏比例确定，移除后的任务进入待修复集合。"
    )

    add_mixed_paragraph(doc, [
        ("text", "9.根据权利要求8所述的一种基于ALNS的无人机无人车协同绿色配送优化方法，其特征在于，所述ALNS算法中的修复算子至少包括贪婪插入算子、后悔值插入算子和中继感知修复算子；所述破坏算子和修复算子的选择权重根据搜索贡献度自适应更新，其更新关系为"),
        ("math", [sup("w", "new"), mr("=(1-ρ)"), sup("w", "old"), mr("+ρs")]),
        ("text", "，其中，"),
        ("math", [mr("ρ")]),
        ("text", "表示权重更新系数，"),
        ("math", [mr("s")]),
        ("text", "表示算子在当前统计周期内获得的贡献评分；当新解优于当前解或优于历史最优解时，提高对应算子的后续选择概率。"),
    ])

    add_text_paragraph(
        doc,
        "10.根据权利要求1所述的一种基于ALNS的无人机无人车协同绿色配送优化方法，其特征在于，所述步骤6中的动态重规划与回退处理机制具体包括以下子步骤：步骤61，检测中继节点有效性、无人机剩余电量、无人车到达状态和任务时间窗剩余量；步骤62，当原接力配送方案仍可通过更换中继节点或更换无人车恢复时，优先执行局部重规划；步骤63，当接力配送方案不可恢复但直送方案满足约束时，将任务切换为无人机直送方式；步骤64，当直送方案仍不可行时，将该任务恢复为待分配状态，并在下一轮ALNS迭代中重新参与候选构建、成本评估和设备匹配；步骤65，记录重规划次数、回退次数、不可行任务数量和最终方案成本，并将其作为输出配送方案的评价信息。"
    )

    doc.save(OUTPUT_DOC)


if __name__ == "__main__":
    main()
