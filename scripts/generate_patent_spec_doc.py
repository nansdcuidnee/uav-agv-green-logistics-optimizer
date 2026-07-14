from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


TEMPLATE_DOC = Path(
    r"D:\师兄们的专利\专利_一种基于可微分的表面缺陷检测方法\专利_一种基于可微分的表面缺陷检测方法\说明书 .docx"
)
OUTPUT_DOC = Path(
    r"D:\uav-agv-green-logistics-optimizer\docs\说明书_一种基于ALNS的无人机无人车协同绿色配送优化方法.docx"
)


def clear_document(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_run_font(run, size: int = 12, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        rpr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "宋体")


def add_text_paragraph(
    doc: Document,
    text: str,
    *,
    center: bool = False,
    bold: bool = False,
) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, 12, bold)


def mr(text: str):
    run = OxmlElement("m:r")
    text_node = OxmlElement("m:t")
    text_node.text = text
    run.append(text_node)
    return run


def math(nodes):
    omath = OxmlElement("m:oMath")
    for node in nodes:
        omath.append(node)
    return omath


def sub(base: str, subscript: str):
    node = OxmlElement("m:sSub")
    e = OxmlElement("m:e")
    e.append(mr(base))
    s = OxmlElement("m:sub")
    s.append(mr(subscript))
    node.append(e)
    node.append(s)
    return node


def sup(base: str, superscript: str):
    node = OxmlElement("m:sSup")
    e = OxmlElement("m:e")
    e.append(mr(base))
    s = OxmlElement("m:sup")
    s.append(mr(superscript))
    node.append(e)
    node.append(s)
    return node


def frac(num: str, den: str):
    node = OxmlElement("m:f")
    num_node = OxmlElement("m:num")
    den_node = OxmlElement("m:den")
    num_node.append(mr(num))
    den_node.append(mr(den))
    node.append(num_node)
    node.append(den_node)
    return node


def add_math_paragraph(doc: Document, nodes) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    para = OxmlElement("m:oMathPara")
    para.append(math(nodes))
    p._element.append(para)


def add_mixed_paragraph(doc: Document, parts) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    for kind, value in parts:
        if kind == "text":
            run = p.add_run(value)
            set_run_font(run)
        elif kind == "math":
            p._p.append(math(value))
        else:
            raise ValueError(kind)


def build_doc() -> Document:
    doc = Document(str(TEMPLATE_DOC))
    clear_document(doc)

    if "Normal" in doc.styles:
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(12)
        doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    add_text_paragraph(doc, "一种基于ALNS的无人机无人车协同绿色配送优化方法", center=True, bold=True)

    add_text_paragraph(doc, "技术领域")
    add_text_paragraph(
        doc,
        "[0001] 本发明涉及低空物流配送与协同调度优化技术领域，尤其涉及一种基于自适应大邻域搜索算法（ALNS）的无人机与无人车协同绿色配送优化方法。",
    )

    add_text_paragraph(doc, "背景技术")
    add_text_paragraph(
        doc,
        "[0002] 随着即时配送、园区物流和校园物流的发展，单纯依靠无人机执行末端配送容易受到续航里程、剩余电量、载荷能力和起降位置的限制，单纯依靠地面车辆执行配送又存在绕行距离长、时效不足的问题，因此空地协同配送逐渐成为提升配送效率和降低单位任务能耗的重要方向。",
    )
    add_text_paragraph(
        doc,
        "[0003] 现有空地协同方案中，常见做法是先为任务预设固定中继点或仅依据静态距离进行模式选择，未能将无人机当前位置、无人车当前位置、接力等待时间、剩余续航裕度以及任务插入位置变化统一纳入同一评价框架，从而导致评分结果与实际执行过程之间存在偏差。",
    )
    add_text_paragraph(
        doc,
        "[0004] 特别是在接力配送场景中，若仅在评分阶段假设无人机已部署到中继点，而执行阶段并未对无人机飞往中继点的轨迹、时间和能耗进行连续建模，则会造成接力方案成本低估、等待风险不可解释以及方案可追溯性不足，难以支撑稳定的协同调度优化。",
    )

    add_text_paragraph(doc, "发明内容")
    add_text_paragraph(
        doc,
        "[0005] 本发明的目的在于提供一种基于ALNS的无人机无人车协同绿色配送优化方法，以解决现有协同配送方案中任务模式选择、接力点选择、评分与执行不一致以及异常情况下回退决策不透明的问题。",
    )
    add_text_paragraph(doc, "[0006] 本发明通过以下技术方案实现：")
    add_text_paragraph(
        doc,
        "[0007] 一种基于ALNS的无人机无人车协同绿色配送优化方法，包括以下步骤：",
    )
    add_text_paragraph(
        doc,
        "[0008] 步骤一：获取配送任务集合、无人机集合、无人车集合以及配送环境约束信息；其中，所述配送任务至少包括任务起点、任务终点、载荷、优先级和截止时间信息，所述无人机至少包括当前位置、剩余电量、最大航程和飞行速度信息，所述无人车至少包括当前位置和行驶状态信息。",
    )
    add_text_paragraph(
        doc,
        "[0009] 步骤二：针对每个任务与无人机组合构建候选服务方式集合，所述候选服务方式至少包括无人机直接配送方式和无人机与无人车接力配送方式；其中接力配送方式的候选中继点由无人车当前位置、任务起点、任务终点、任务走廊采样点以及无人车到任务走廊的投影点生成。",
    )
    add_text_paragraph(
        doc,
        "[0010] 步骤三：对所述候选服务方式进行可行性筛选，筛除不满足位置有效性、剩余续航、配送完成时间和接力部署条件的候选方案；其中接力配送方式要求无人机能够从当前位置连续飞行至候选中继点，并在到达中继点后仍具备完成配送段的剩余航程。",
    )
    add_text_paragraph(
        doc,
        "[0011] 步骤四：构建统一评价模型，对直接配送方式和接力配送方式分别计算时间成本、无人机能耗、无人车能耗、等待惩罚和回退风险，得到对应候选方案的综合成本。",
    )
    add_text_paragraph(
        doc,
        "[0012] 步骤五：基于步骤四得到的候选方案成本，采用启发式插入方式生成初始解，并通过ALNS中的破坏算子与修复算子对当前解进行迭代优化，获得任务分配、设备匹配、路线插入位置和中继点选择的联合优化结果。",
    )
    add_text_paragraph(
        doc,
        "[0013] 步骤六：在执行过程中，当接力方案出现部署不可行、同步偏差过大、等待超时或剩余裕度不足时，触发基于同一评价器的回退或重选机制，将当前任务切换为新的接力方案或直接配送方案。",
    )
    add_text_paragraph(
        doc,
        "[0014] 步骤七：输出满足约束条件的无人机无人车协同配送方案，所述配送方案至少包括任务分配结果、直接配送或接力配送模式、中继点位置、设备执行顺序以及能耗与时序评估结果。",
    )

    add_text_paragraph(doc, "[0015] 进一步地，所述直接配送方式按照无人机当前位置、任务起点、任务终点以及返回锚点形成配送距离，所述接力配送方式按照无人机当前位置、中继点、任务起点、任务终点以及中继点形成配送距离。")
    add_math_paragraph(
        doc,
        [
            sub("D", "direct"),
            mr("="),
            mr("d(u,s)"),
            mr("+"),
            mr("d(s,e)"),
            mr("+"),
            mr("d(e,u)"),
        ],
    )
    add_math_paragraph(
        doc,
        [
            sub("D", "relay"),
            mr("="),
            mr("d(u,r)"),
            mr("+"),
            mr("d(r,s)"),
            mr("+"),
            mr("d(s,e)"),
            mr("+"),
            mr("d(e,r)"),
        ],
    )
    add_mixed_paragraph(
        doc,
        [
            ("text", "[0016] 其中，"),
            ("math", [mr("u")]),
            ("text", "表示无人机当前位置，"),
            ("math", [mr("r")]),
            ("text", "表示中继点，"),
            ("math", [mr("s")]),
            ("text", "表示任务起点，"),
            ("math", [mr("e")]),
            ("text", "表示任务终点，"),
            ("math", [mr("d(a,b)")]),
            ("text", "表示节点a与节点b之间的距离。"),
        ],
    )

    add_text_paragraph(doc, "[0017] 进一步地，所述统一评价模型中的时间成本采用距离与速度的比值计算，无人机能耗与无人车能耗分别采用距离与单位能耗参数的乘积计算。")
    add_math_paragraph(
        doc,
        [
            sub("T", "mode"),
            mr("="),
            frac("D_mode", "v_u"),
        ],
    )
    add_math_paragraph(
        doc,
        [
            sub("E", "u"),
            mr("="),
            frac("D_mode", "1000"),
            mr("·"),
            sub("η", "u"),
        ],
    )
    add_math_paragraph(
        doc,
        [
            sub("E", "a"),
            mr("="),
            frac("d(a,r)", "1000"),
            mr("·"),
            sub("η", "a"),
        ],
    )
    add_mixed_paragraph(
        doc,
        [
            ("text", "[0018] 其中，"),
            ("math", [sub("T", "mode")]),
            ("text", "表示候选服务方式的时间成本，"),
            ("math", [mr("D_mode")]),
            ("text", "表示对应服务方式的飞行距离，"),
            ("math", [mr("v_u")]),
            ("text", "表示无人机飞行速度，"),
            ("math", [sub("η", "u")]),
            ("text", "表示无人机单位距离能耗参数，"),
            ("math", [sub("η", "a")]),
            ("text", "表示无人车单位距离能耗参数。"),
        ],
    )

    add_text_paragraph(doc, "[0019] 进一步地，接力配送方式中的等待惩罚基于无人车到达中继点的预计时刻与无人机到达中继点的预计时刻之间的时差确定。")
    add_math_paragraph(
        doc,
        [
            sub("W", "relay"),
            mr("="),
            mr("|"),
            sub("t", "a"),
            mr("-"),
            sub("t", "u"),
            mr("|"),
        ],
    )
    add_text_paragraph(
        doc,
        "[0020] 其中，ta表示无人车到达中继点的预计时刻，tu表示无人机从当前位置连续飞行至中继点的预计到达时刻，所述等待惩罚与所述时差成正相关，用于表征接力配送中的同步偏差。",
    )

    add_text_paragraph(doc, "[0021] 进一步地，所述综合成本采用加权求和方式构建：")
    add_math_paragraph(
        doc,
        [
            mr("C"),
            mr("="),
            sub("w", "t"),
            sub("T", "mode"),
            mr("+"),
            sub("w", "u"),
            sub("E", "u"),
            mr("+"),
            sub("w", "a"),
            sub("E", "a"),
            mr("+"),
            sub("w", "w"),
            sub("W", "relay"),
            mr("+"),
            sub("w", "f"),
            sup("R", "risk"),
        ],
    )
    add_text_paragraph(
        doc,
        "[0022] 其中，wt、wu、wa、ww和wf分别为时间成本、无人机能耗、无人车能耗、等待惩罚和回退风险的权重系数，Rrisk表示由剩余续航裕度、时间裕度和同步偏差共同派生的风险项。",
    )

    add_text_paragraph(
        doc,
        "[0023] 本发明相对于现有技术具有如下有益效果：第一，将无人机直接配送与接力配送统一纳入同一评价框架，便于在任务分配与插入决策中进行一致比较；第二，在接力配送中显式计入无人机从当前位置连续飞行到中继点的部署距离、部署时间和部署能耗，使评分结果与执行过程保持一致；第三，通过将中继点可行性、等待偏差和回退风险纳入统一评价器，提高了异常情况下方案切换的可解释性；第四，利用ALNS对任务插入位置和服务模式进行联合优化，能够在满足约束条件的前提下降低单位任务能耗并改善任务完成效率。",
    )

    add_text_paragraph(doc, "附图说明")
    add_text_paragraph(doc, "[0024] 图1为本发明一种基于ALNS的无人机无人车协同绿色配送优化方法的总体流程示意图；")
    add_text_paragraph(doc, "[0025] 图2为本发明中直接配送方式与接力配送方式的统一候选评估示意图；")
    add_text_paragraph(doc, "[0026] 图3为本发明中接力配送方式的连续部署执行时序示意图；")
    add_text_paragraph(doc, "[0027] 图4为本发明中ALNS破坏与修复迭代优化过程示意图。")

    add_text_paragraph(doc, "具体实施方式")
    add_text_paragraph(
        doc,
        "[0028] 为使本发明的目的、技术方案和有益效果更加清楚，以下结合实施方式对本发明作进一步说明。应当理解，以下实施方式仅用于说明本发明，而不用于限定本发明的保护范围。",
    )
    add_text_paragraph(doc, "[0029] 实施例一：")
    add_text_paragraph(
        doc,
        "[0030] 在本实施例中，构建包含无人机、无人车、配送任务和环境约束的配送场景。任务对象至少包含任务起点、任务终点、载荷、优先级和截止时间；无人机对象至少包含当前位置、电量、最大航程和飞行速度；无人车对象至少包含当前位置和状态信息；环境对象至少包含可通行边界、障碍物或禁行区域信息。",
    )
    add_text_paragraph(
        doc,
        "[0031] 对于每个待分配任务，系统首先以无人机当前位置为真实起点生成直接配送候选方案，并计算该任务插入无人机当前路线不同位置时的增量成本；然后针对空闲无人车生成与中继点绑定的接力配送候选方案，所述中继点包括无人车当前位置、任务起点、任务终点、起终点走廊上的采样点以及无人车到走廊线段的投影点。",
    )
    add_text_paragraph(
        doc,
        "[0032] 在接力配送候选方案中，系统先判断候选中继点是否位于有效位置，再判断无人机从当前位置飞往该中继点后的剩余航程是否足以完成从中继点到任务起点、任务终点并返回中继点的配送段；若任务含有截止时间，则进一步判断总配送时间是否不超过所述截止时间，从而完成候选中继点的预筛选。",
    )
    add_text_paragraph(
        doc,
        "[0033] 对预筛选后的直接配送候选方案和接力配送候选方案，统一调用评价器计算时间成本、无人机能耗、无人车能耗和等待惩罚。对于接力配送方案，还计算基于剩余续航裕度、任务时间裕度和同步偏差的回退风险；评价器输出综合成本以及可行性标志，供后续插入与重规划使用。",
    )
    add_text_paragraph(
        doc,
        "[0034] 在生成初始解时，系统按照任务插入成本从低到高将任务插入无人机或无人车路线；在ALNS迭代过程中，通过破坏算子移除部分已分配任务，通过修复算子结合统一评价器重新选择服务模式、插入位置和中继点，并利用模拟退火准则决定是否接受新解。",
    )
    add_text_paragraph(
        doc,
        "[0035] 在执行阶段，若任务采用接力配送方式，则无人车首先向中继点移动，无人机在无人车到达中继点后由waiting_for_agv状态推进至uav_flying_to_relay状态，并沿规划路径连续飞行至中继点；当无人机到达中继点后触发部署完成事件，再进入配送执行阶段。该过程中无人机部署段的轨迹、时间和能耗被计入任务统计结果。",
    )
    add_text_paragraph(
        doc,
        "[0036] 若执行过程中检测到原接力方案部署不可行、等待偏差超过阈值、时间裕度不足或等待超时，则系统调用与初始评分一致的统一评价器重新评估当前任务的可选方案；当存在更优可行接力方案时，执行中继点重选；当接力方案整体不可行时，回退为直接配送方案，从而保证评分逻辑与异常处理逻辑一致。",
    )
    add_text_paragraph(
        doc,
        "[0037] 本实施例中，统一评价器和执行过程使用相同的接力语义：无人机接力配送总路径包含从当前位置到中继点的部署段，以及从中继点至任务起点、任务终点并回到中继点的配送段；因此，评价成本、可行性筛选结果和执行统计结果之间能够形成一一对应关系，便于后续验证与复盘。",
    )
    add_text_paragraph(
        doc,
        "[0038] 应当说明，本实施例中的时间成本与能耗成本采用简化的距离比例模型，主要用于调度决策评价，而非高保真飞行动力学建模；在不脱离本发明基本思想的前提下，本领域技术人员可根据具体应用场景对权重系数、采样密度或风险阈值进行调整，均应落入本发明的保护范围。",
    )

    return doc


def main() -> None:
    doc = build_doc()
    OUTPUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_DOC))
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
